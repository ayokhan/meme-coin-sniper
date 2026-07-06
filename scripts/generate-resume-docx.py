"""Generate AYO-PERVEZ-KHAN-AI-Product-Manager-Resume.docx from structured content."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "AYO-PERVEZ-KHAN-AI-Product-Manager-Resume-v5.docx"


def add_rich_paragraph(doc: Document, text: str, *, bullet: bool = False, italic: bool = False):
    style = "List Bullet" if bullet else None
    p = doc.add_paragraph(style=style)
    if italic:
        p.paragraph_format.space_after = Pt(2)
    for part in re.split(r"(\*\*.*?\*\*)", text):
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
            run.italic = italic
        else:
            run = p.add_run(part)
            run.italic = italic
    return p


def add_section_heading(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_job_heading(doc: Document, title: str, subtitle: str):
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    add_rich_paragraph(doc, subtitle, italic=True)


def build() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    # Name
    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name.add_run("AYO PERVEZ KHAN")
    run.bold = True
    run.font.size = Pt(18)

    # Tagline
    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tag.add_run(
        "Senior AI Product Manager · Product Owner · PMP® · CSPO®"
    )
    run.bold = True
    run.font.size = Pt(11)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run(
        "Innisfil, ON · ayokhan2006@gmail.com · 905-904-3654 · "
    )
    link = contact.add_run("LinkedIn")
    link.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
    link.underline = True

    add_section_heading(doc, "Summary")
    add_rich_paragraph(
        doc,
        "Product leader with **14+ years** shipping enterprise and consumer platforms. At "
        "**Enercare** (**1.6M+ customers**), I own **COMPASS** (shipped **Launch 1**—unified "
        "retention desktop for **110+ agents**, **RWH/RHVAC**), led **Project Billy** "
        "(**Enbridge → Enercare**, **Zuora Billing**), **Project Maple** CRM, **AI voice/chat "
        "agents**, and **My Account** (0→1 self-service for residential owners, landlord/tenants, "
        "builders, commercial). At **NovaStaris**, I practice **AI product management** end-to-end—"
        "**JTBD**, **Figma/Miro** prototypes, **RAG** and **evals**, stakeholder-led releases on "
        "**Vercel**. Seeking **AI PM, PM, or PO** roles.",
    )

    add_section_heading(doc, "Signature Impact")
    for item in [
        "**0→1 & scale:** **NovaStaris**—AI SaaS (web + Play Store) with subscriptions and "
        "feature-flag rollouts; **Enercare**—**COMPASS Launch 1** (unified retention desktop for "
        "**110+ agents daily**, **RWH/RHVAC**), **1.6M+ customers**, **Dynamics 365, Salesforce, Zuora**.",
        "**Production AI:** **NovaStaris**—**RAG**, **eval loops**, prompt specs, and **good/bad "
        "feedback** in production; **Enercare**—**Vapi** voice and **N8N** chat agents replacing "
        "legacy call flows.",
        "**Roadmap & high-stakes delivery:** **Enercare**—**Project Billy**, **Project Maple**; "
        "**NovaStaris**—VIP roadmap, **JTBD**-led features, Google Play closed testing, phased "
        "releases via **feature flags**.",
        "**Agile at scale:** Lead **22-person** distributed delivery (Enercare) in **JIRA**; "
        "**NovaStaris**—backlog refinement, **GitHub**, and **Vercel** production releases.",
    ]:
        add_rich_paragraph(doc, item, bullet=True)

    add_section_heading(doc, "Professional Experience")

    add_job_heading(
        doc,
        "Founder & AI Product Manager — NovaStaris (novastaris.ai)",
        "*2024 – Present · AI crypto trading intelligence SaaS*",
    )
    for item in [
        "**JTBD, vision & roadmap:** Framed **jobs-to-be-done** for meme discovery, crypto futures, "
        "prediction markets, and VIP workflows; prioritized roadmap across Pro/VIP tiers, on-demand "
        "access, and mobile—balancing revenue, AI differentiation, and platform constraints.",
        "**Business problems solved:** Retail traders juggled **multiple tools** for meme discovery, "
        "wallet tracking, futures/perps, and AI chart reads—slow manual research, missed early "
        "entries, and **no reliable loop** to improve AI output quality. Built a unified **AI "
        "trading intelligence** workspace on **novastaris.ai** (Solana/BSC discovery, wallet "
        "trackers, **Perp Radar**, liquidation map, prediction-market tools) with **RAG**, "
        "production **evals** and **good/bad feedback**, in-app **early breakout alerts** "
        "(**NovaPick** / **LATE CHASE** for scalps), and **Pro/VIP** subscriptions (web + "
        "**Google Play**).",
        "**Discovery & prototyping:** Built **Figma** and **Miro** prototypes for AI analysis, Nova "
        "Radar, subscriptions, and admin flows; validated concepts with **JTBD** framing and early "
        "subscriber feedback before release.",
        "**AI PM — specs & quality:** Authored PRDs and acceptance criteria for **AI Analysis**, "
        "**Nova Radar**, **Nova Forex Agent**, **per-user RAG**, and offline **eval** tooling; "
        "defined prompt structure, output schema, **good/bad feedback** capture, and regression "
        "metrics from production data.",
        "**Stakeholders & release planning:** Aligned **VIP subscribers**, support, and GTM on "
        "phased launches—**feature-flag** rollouts, **Vercel** production releases, **Google Play** "
        "closed testing, and **Stripe** billing policy for web vs mobile access.",
        "**Shipped AI product surface:** Launched token scoring (Solana/BSC), **Crypto Futures**, "
        "**Liquidation Map**, wallet trackers, prediction-market tools, and Telegram/coach signals—"
        "each scoped as an MVP with clear success criteria.",
        "**Backlog & prioritization:** Owned product backlog and refinement; sized epics and wrote "
        "**user stories** by business value—sprint-ready specs with acceptance criteria and release notes.",
        "**Monetization & analytics:** Owned **Stripe** subscription tiers, admin customer hub, and "
        "**PostHog** product analytics; informed roadmap from conversion, retention, and feature usage.",
        "**AI-augmented PM stack:** **Cursor**, **GitHub**, and **Vercel** for rapid PRD-to-production "
        "iteration; **Claude** and **OpenAI** APIs for LLM features, embeddings, and eval pipelines.",
    ]:
        add_rich_paragraph(doc, item, bullet=True)

    add_job_heading(
        doc,
        "Product Owner & QA Team Lead — Enercare Inc.",
        "*March 2021 – Present · Enterprise utilities & CRM · Product Manager, COMPASS*",
    )
    for item in [
        "**COMPASS (Product Manager):** Own **COMPASS**—call-center platform for **RWH/RHVAC "
        "retention offers**; **Launch 1** live **March 2026** (**8 months** kickoff→production), "
        "**110+ agents** on COMPASS daily. Roadmap, backlog, and releases for offer types, "
        "eligibility rules, and agent workflow.",
        "**COMPASS Launch 1 — business problems solved:** Before launch, agents worked across "
        "**Salesforce, Clarify, and Zuora** (\"swivel chair\")—manual offer math caused "
        "**calculation errors** flowing to billing/finance; a mandated **back-office review queue** "
        "delayed promised adjustments in peak season; customers saw a **promise-vs-bill gap**; new "
        "offers took **10+ weeks** to ship; **manual case creation** drove handle time up and "
        "quality down.",
        "**COMPASS Launch 1 — what we shipped:** Unified **Customer 360 agent desktop** (one screen, "
        "one truth); **Discount Rules Engine** shifted offer logic from code to **configuration** "
        "(eligibility rules, effective/expiry dates—**months/weeks → days**); modular "
        "**Orchestrator + CDS** spine with direct **Zuora** integration (system of record) and "
        "**Clarify** back-sync; **~12.5K** cases auto-created post-launch.",
        "**COMPASS Launch 1 — outcomes:** **~17K** interactions processed; **~16.9K** successful "
        "downstream **Clarify & Zuora** calls; **0 P1/P2** production tickets; strong **Day 1 "
        "adoption** with sharp drops in manual case creation, cross-system navigation, and "
        "calculation errors; **8** scope additions delivered above baseline.",
        "**Project Billy (Product Manager):** Led product for **Project Billy**—**Enbridge → Enercare** "
        "billing migration of **1.6 million customers** onto **Zuora Billing**; PRD, cutover planning, "
        "billing/regulatory alignment with **Enbridge**, UAT, and hypercare with operations and "
        "finance stakeholders.",
        "**My Account (0→1):** Drive **My Account**—Enercare's customer self-service product for "
        "**residential owners, landlord/tenants, builders, and commercial** segments—from competitive "
        "research, ICP, **JTBD**, and **Figma** prototype through functional requirements and UAT.",
        "**Vision, roadmap & growth:** Own product vision and roadmap across Enercare's **1.6M+ "
        "customer** base—prioritized **COMPASS** retention, **My Account**, **Zuora** billing "
        "capabilities, **AI agent** strategy, and **Project Maple** CRM modernization.",
        "**MVPs with cross-functional teams:** Delivered **COMPASS** retention-offer enhancements "
        "for **RWH/RHVAC**, **My Account** MVP by customer segment, **Vapi voice** and **N8N chat** "
        "agents (PRD → pilot → production), and releases on **Dynamics 365 / Salesforce**.",
        "**Backlog, refinement & epic sizing:** Lead **refinement** and **epic sizing** in **JIRA** "
        "for multi-system backlog; decomposed **Project Billy** and **Maple** CRM work into "
        "cutover-ready stories with NFRs for a **22-member** squad (Canada, India, USA).",
        "**Release planning & stakeholders:** Plan releases and set delivery expectations with "
        "**operations**, **Enbridge**, and vendors—**Billy** and **Maple** cutover milestones, "
        "regression/UAT, rollback criteria, and weekly sponsor updates.",
        "**High-stakes migration:** Led **Project Maple** (**Salesforce → Clarify CRM**) end-to-end "
        "under fixed Enbridge deadline—discovery, data mapping, cutover runbooks, hypercare, and "
        "sign-off.",
        "**AI agents in production:** Authored PRDs and quality bars for **voice (Vapi)** and "
        "**chat (N8N)** agents; replaced underperforming legacy AI call flows serving **~1/5 of "
        "Ontario homes** with measurable containment and escalation paths.",
        "**Enable focused sprints:** Run sprint planning with clear sprint goals; ensure stories are "
        "**ready to pick up** (dependencies, test data, env access) before sprint start—reduced "
        "carry-over on CRM and billing epics.",
        "**AI-augmented PM:** Use **Claude, Cursor, Figma AI** for PRDs, journey maps, and faster "
        "iteration with onshore/offshore engineering.",
    ]:
        add_rich_paragraph(doc, item, bullet=True)

    add_job_heading(
        doc,
        "Product Manager / QA Lead — Mackenzie Health",
        "*June 2020 – March 2021 · SMART Hospital (first in Canada)*",
    )
    for item in [
        "**Team leadership:** Led **QA** and **business analyst** teams for Canada's first "
        "**SMART Hospital**—authored **requirements** and **user stories**, owned test strategy, "
        "and coordinated delivery across **6 mission-critical systems** (Epic, Vocera, RTLS, nurse "
        "call, patient engagement).",
        "**Business problems solved:** Patients lacked **bedside access to education and care** from "
        "the hospital bed; **caregiver–provider communication** was slow and fragmented; vulnerable "
        "populations (**mental health**, **newborns**, **mother–child**) needed **real-time "
        "monitoring and safety tracking**. Delivered integrated SMART Hospital capabilities—bedside "
        "**patient education and engagement**, **Vocera** and nurse-call workflows for efficient "
        "staff communication, and **RTLS** location tracking for at-risk patients.",
        "**HL7 integration testing:** Led **HL7** interface testing and validation between clinical "
        "and operational systems; brokered multi-vendor integration specs with US partners and "
        "translated workflows into release-ready acceptance criteria.",
        "**Pre-launch validation:** Ran super-user research, training, and clinical workflow sign-off "
        "for smart bed, RTLS, and nurse call capabilities ahead of hospital go-live.",
    ]:
        add_rich_paragraph(doc, item, bullet=True)

    add_job_heading(
        doc,
        "Product / Program Lead — CB2 Insights (Skylight Health Group)",
        "*October 2018 – May 2020 · Healthcare SaaS*",
    )
    for item in [
        "Owned product strategy and QA for **6 concurrent SaaS platforms** (200+ providers): "
        "roadmaps, backlog prioritization, and release quality.",
        "Delivered CI/CD automation that **cut deployment cycles ~70%**, enabling faster feature iteration.",
        "Evangelized Agile (JIRA, Confluence); API validation (Postman/JMeter) for third-party integrations.",
    ]:
        add_rich_paragraph(doc, item, bullet=True)

    add_job_heading(
        doc,
        "QA / Business Analyst — Cancer Care Ontario",
        "*February 2012 – October 2018 · Provincial healthcare*",
    )
    for item in [
        "Led QA strategy for Ontario's largest integration programs (Wait Times, Cardiac, Renal, "
        "Emergency Triage)—systems serving **14M residents**.",
        "Drove **Waterfall → Agile** transformation; mentored analysts; owned BRDs, test strategy, "
        "and vendor triage for HL7/BizTalk integrations.",
    ]:
        add_rich_paragraph(doc, item, bullet=True)

    add_section_heading(doc, "Education")
    for item in [
        "**MBA, Information Technology** — University of Cumbria",
        "**Advanced Diploma, Social Works** — University of Lagos",
        "**Project Management Specialization** — University of California, Irvine",
    ]:
        add_rich_paragraph(doc, item, bullet=True)

    add_section_heading(doc, "Core Competencies")
    for item in [
        "**Product:** Product vision · Roadmap (business value + tech tradeoffs) · **MVP definition** · "
        "PRD · **User stories & acceptance criteria** · **Backlog grooming, epic sizing & refinement** · "
        "**Release planning & sprint goals** · GTM & pricing · **Stakeholder communication** · UAT & launch · "
        "Feature flags · Subscription SaaS",
        "**AI / LLM:** Agent design (voice, chat, workflow) · Prompt engineering · Evals & feedback "
        "loops · RAG · Fine-tuning strategy · Guardrails & responsible AI · AI quality metrics",
        "**Tools:** Claude · **Cursor** · OpenAI API · **Figma** · **Miro** · **PostHog** · **GitHub** · "
        "**Vercel** · N8N · Vapi · JIRA · Confluence · Postman · Azure DevOps",
        "**Domains:** Fintech/crypto intelligence · Utilities & CRM · Healthcare (Epic, HL7) · "
        "Enterprise agile at scale",
    ]:
        add_rich_paragraph(doc, item, bullet=True)

    add_section_heading(doc, "Certifications")
    doc.add_paragraph(
        "PMP® · CSPO® (Scrum Alliance) · CTFL® (ISTQB) · Scrum Master Accredited · "
        "Agile Scrum Foundation (EXIN) · Business Analyst RPA (UiPath) · "
        "Accredited Software Test Manager"
    )

    foot = doc.add_paragraph()
    foot.paragraph_format.space_before = Pt(8)
    run = foot.add_run(
        "NovaStaris is an independent product initiative alongside full-time role at Enercare Inc."
    )
    run.italic = True
    run.font.size = Pt(9)

    return doc


if __name__ == "__main__":
    build().save(OUTPUT)
    print(f"Wrote {OUTPUT}")
