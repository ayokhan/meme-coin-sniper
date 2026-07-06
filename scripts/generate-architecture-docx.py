"""Generate NovaStaris-Implementation-Architecture.docx from docs/NOVASTARIS_IMPLEMENTATION_ARCHITECTURE.md"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "NOVASTARIS_IMPLEMENTATION_ARCHITECTURE.md"
OUTPUT = ROOT / "docs" / "NovaStaris-Implementation-Architecture.docx"


def add_rich_runs(paragraph, text: str, *, monospace: bool = False):
    for part in re.split(r"(\*\*.*?\*\*|`[^`]+`)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        else:
            run = paragraph.add_run(part)
        if monospace:
            run.font.name = "Consolas"
            run.font.size = Pt(8.5)


def set_cell_shading(cell, fill: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci in range(col_count):
            cell = table.rows[ri].cells[ci]
            text = row[ci] if ci < len(row) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            add_rich_runs(p, text.strip())
            for run in p.runs:
                run.font.size = Pt(9)
            if ri == 0:
                set_cell_shading(cell, "E8EEF4")
                for run in p.runs:
                    run.bold = True
    doc.add_paragraph()


def parse_markdown_table(lines: list[str]) -> list[list[str]] | None:
    if len(lines) < 2:
        return None
    if not all("|" in ln for ln in lines[:2]):
        return None
    if not re.match(r"^\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$", lines[1].strip()):
        return None

    def split_row(line: str) -> list[str]:
        line = line.strip().strip("|")
        return [c.strip() for c in line.split("|")]

    return [split_row(ln) for ln in lines if ln.strip() and not re.match(r"^\|?\s*:?-+", ln.strip())]


def build_from_markdown(md_path: Path) -> Document:
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)

    i = 0
    in_code = False
    code_lines: list[str] = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.15)
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(8)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if stripped == "---":
            doc.add_paragraph()
            i += 1
            continue

        if stripped.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(stripped[2:].strip())
            run.bold = True
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
            i += 1
            continue

        if stripped.startswith("## "):
            p = doc.add_paragraph()
            run = p.add_run(stripped[3:].strip())
            run.bold = True
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x2D, 0x3A, 0x8C)
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        if stripped.startswith("### "):
            p = doc.add_paragraph()
            run = p.add_run(stripped[4:].strip())
            run.bold = True
            run.font.size = Pt(11)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            rows = parse_markdown_table(table_lines)
            if rows:
                add_table(doc, rows)
            continue

        if stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            add_rich_runs(p, stripped[2:].strip(), monospace=False)
            for run in p.runs:
                run.italic = True
            i += 1
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_rich_runs(p, stripped[2:].strip())
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        p = doc.add_paragraph()
        add_rich_runs(p, stripped)
        i += 1

    return doc


if __name__ == "__main__":
    import sys

    source = Path(sys.argv[1]) if len(sys.argv) > 1 else SOURCE
    output = Path(sys.argv[2]) if len(sys.argv) > 2 else OUTPUT
    if not source.is_absolute():
        source = ROOT / source
    if not output.is_absolute():
        output = ROOT / output
    if not source.exists():
        raise SystemExit(f"Source not found: {source}")
    build_from_markdown(source).save(output)
    print(f"Wrote {output}")
